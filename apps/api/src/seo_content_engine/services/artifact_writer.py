from __future__ import annotations

import html
import json
import logging
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt

from seo_content_engine.core.config import settings
from seo_content_engine.services.output_formatter import OutputFormatter
from seo_content_engine.services.schema_markup_generator import SchemaMarkupGenerator
from seo_content_engine.utils.formatters import slugify

logger = logging.getLogger(__name__)


class ArtifactWriter:
    # Map raw data field names to human-readable table column headers shown to buyers.
    COLUMN_DISPLAY_NAMES: dict[str, str] = {
        "quarterName": "Quarter",
        "locationRate": "Locality Rate (₹/sq ft)",
        "micromarketRate": "Micromarket Rate (₹/sq ft)",
        "cityRate": "City Rate (₹/sq ft)",
        "key": "Configuration",
        "doc_count": "Listings",
        "name": "Area",
        "distance_km": "Distance (km)",
        "sale_count": "Resale Listings",
        "sale_avg_price_per_sqft": "Avg Sale Price (₹/sq ft)",
        "url": "Link",
        "avgRate": "Avg Rate (₹/sq ft)",
        "changePercentage": "Change (%)",
        "propertyType": "Property Type",
        "avgPrice": "Avg Sale Price (₹/sq ft)",
        "changePercent": "Change (%)",
        "status": "Project Status",
        "units": "Units",
        "total_listings": "Total Listings",
        "total_projects": "Total Projects",
    }

    @staticmethod
    def _column_label(column_name: str) -> str:
        """Return the human-readable display label for a data column name."""
        return ArtifactWriter.COLUMN_DISPLAY_NAMES.get(column_name, column_name)

    @staticmethod
    def _ensure_artifacts_dir() -> Path:
        artifacts_dir = Path(settings.artifacts_dir)
        artifacts_dir.mkdir(parents=True, exist_ok=True)
        return artifacts_dir

    @staticmethod
    def write_json_artifact(payload: dict, file_stem: str) -> str:
        artifacts_dir = ArtifactWriter._ensure_artifacts_dir()
        output_path = artifacts_dir / f"{slugify(file_stem)}.json"

        with output_path.open("w", encoding="utf-8") as file:
            json.dump(payload, file, ensure_ascii=False, indent=2)

        return str(output_path)

    @staticmethod
    def write_markdown_artifact(markdown_text: str, file_stem: str) -> str:
        artifacts_dir = ArtifactWriter._ensure_artifacts_dir()
        output_path = artifacts_dir / f"{slugify(file_stem)}.md"

        with output_path.open("w", encoding="utf-8") as file:
            file.write(markdown_text)

        return str(output_path)

    @staticmethod
    def _guard_review_block(draft: dict) -> None:
        # Export is always allowed. When needs_review is True (hard validation fail),
        # we log a warning so the exported artifact is clearly flagged but never blocked.
        if draft.get("needs_review"):
            entity_name = draft.get("entity", {}).get("entity_name", "unknown")
            approval = draft.get("quality_report", {}).get("approval_status", "unknown")
            logger.warning(
                "Exporting draft with needs_review=True for entity=%s approval_status=%s. "
                "Artifact will include review flag.",
                entity_name,
                approval,
            )

    @staticmethod
    def _set_base_doc_styles(document: Document) -> None:
        styles = document.styles
        if "Normal" in styles:
            styles["Normal"].font.name = "Arial"
            styles["Normal"].font.size = Pt(10.5)

    @staticmethod
    def _safe_text(value: Any) -> str:
        if value is None:
            return "—"
        return str(value).strip() or "—"

    @staticmethod
    def _html_escape(value: Any) -> str:
        return html.escape(ArtifactWriter._safe_text(value))

    @staticmethod
    def _add_metadata(document: Document, metadata: dict) -> None:
        h1 = ArtifactWriter._safe_text(metadata.get("h1"))
        title = ArtifactWriter._safe_text(metadata.get("title"))
        meta_description = ArtifactWriter._safe_text(metadata.get("meta_description"))
        intro_snippet = ArtifactWriter._safe_text(metadata.get("intro_snippet"))

        document.add_heading(h1, level=1)

        intro_para = document.add_paragraph()
        intro_para.add_run(intro_snippet)

        document.add_heading("SEO Metadata", level=2)
        document.add_paragraph(f"Title: {title}")
        document.add_paragraph(f"Meta Description: {meta_description}")

    @staticmethod
    def _add_sections(document: Document, sections: list[dict]) -> None:
        if not sections:
            return

        document.add_heading("Editorial Sections", level=2)
        for section in sections:
            title = ArtifactWriter._safe_text(section.get("title"))
            body = str(section.get("body") or "").strip()
            key_points: list[str] = [
                str(point).strip()
                for point in (section.get("key_points") or [])
                if point and str(point).strip()
            ]

            document.add_heading(title, level=3)
            if not body:
                document.add_paragraph("No grounded narrative available for this section.")
            else:
                for paragraph in body.split("\n"):
                    paragraph = paragraph.strip()
                    if paragraph:
                        document.add_paragraph(paragraph)

            for point in key_points:
                document.add_paragraph(point, style="List Bullet")

    @staticmethod
    def _add_tables(document: Document, tables: list[dict]) -> None:
        if not tables:
            return

        document.add_heading("Key Data Tables", level=2)
        for table_block in tables:
            title = ArtifactWriter._safe_text(table_block.get("title"))
            summary = str(table_block.get("summary") or "").strip()
            columns = table_block.get("columns", []) or []
            rows = table_block.get("rows", []) or []

            document.add_heading(title, level=3)
            if summary:
                document.add_paragraph(summary)

            if not columns:
                document.add_paragraph("No structured columns available.")
                continue

            doc_table = document.add_table(rows=1, cols=len(columns))
            doc_table.style = "Table Grid"
            header_cells = doc_table.rows[0].cells
            for index, column_name in enumerate(columns):
                header_cells[index].text = ArtifactWriter._column_label(column_name)

            if rows:
                for row in rows:
                    cells = doc_table.add_row().cells
                    for index, column_name in enumerate(columns):
                        cells[index].text = OutputFormatter.format_cell(
                            column_name,
                            row.get(column_name),
                        )
            else:
                cells = doc_table.add_row().cells
                for index in range(len(columns)):
                    cells[index].text = "—"

            document.add_paragraph("")

    @staticmethod
    def _add_faqs(document: Document, faqs: list[dict]) -> None:
        if not faqs:
            return

        document.add_heading("Frequently Asked Questions", level=2)
        for faq in faqs:
            question = ArtifactWriter._safe_text(faq.get("question"))
            answer = str(faq.get("answer") or "").strip()

            document.add_heading(question, level=3)
            if not answer:
                document.add_paragraph("No grounded answer available.")
                continue

            for paragraph in answer.split("\n"):
                paragraph = paragraph.strip()
                if paragraph:
                    document.add_paragraph(paragraph)

    @staticmethod
    def _normalize_link_items(links: list[Any]) -> list[dict]:
        normalized_links: list[dict] = []

        for item in links or []:
            if isinstance(item, list):
                for nested in item:
                    if isinstance(nested, dict):
                        normalized_links.append(nested)
            elif isinstance(item, dict):
                normalized_links.append(item)

        return normalized_links

    @staticmethod
    def _add_internal_links(document: Document, internal_links: dict) -> None:
        if not internal_links:
            return

        title_map = {
            "sale_unit_type_links": "Unit Type Links",
            "sale_property_type_links": "Property Type Links",
            "sale_quick_links": "Quick Links",
            "nearby_locality_links": "Nearby Locality Links",
            "top_project_links": "Top Project Links",
            "featured_project_links": "Featured Project Links",
        }

        document.add_heading("Internal Links", level=2)

        for key, group_title in title_map.items():
            group_links = ArtifactWriter._normalize_link_items(internal_links.get(key, []))
            if not group_links:
                continue

            document.add_heading(group_title, level=3)
            for item in group_links:
                label = ArtifactWriter._safe_text(
                    item.get("label")
                    or item.get("unitType")
                    or item.get("propertyType")
                    or "Link"
                )
                url = ArtifactWriter._safe_text(item.get("url"))
                if url != "—":
                    document.add_paragraph(f"{label}: {url}", style="List Bullet")
                else:
                    document.add_paragraph(label, style="List Bullet")

    @staticmethod
    def write_docx_artifact(draft: dict, file_stem: str) -> str:
        artifacts_dir = ArtifactWriter._ensure_artifacts_dir()
        output_path = artifacts_dir / f"{slugify(file_stem)}.docx"

        document = Document()
        ArtifactWriter._set_base_doc_styles(document)

        # E3: Prepend a review banner when the draft failed hard validation.
        if draft.get("needs_review"):
            approval = draft.get("quality_report", {}).get("approval_status", "unknown")
            banner = document.add_paragraph(
                f"⚠️ REVIEW REQUIRED — Validation failed (approval_status={approval}). "
                "Do not publish without editorial review."
            )
            banner.runs[0].bold = True
            from docx.shared import RGBColor
            banner.runs[0].font.color.rgb = RGBColor(0xC0, 0x3A, 0x00)
            document.add_paragraph("")

        ArtifactWriter._add_metadata(document, draft.get("metadata", {}))
        ArtifactWriter._add_sections(document, draft.get("sections", []))
        ArtifactWriter._add_tables(document, draft.get("tables", []))
        ArtifactWriter._add_faqs(document, draft.get("faqs", []))
        ArtifactWriter._add_internal_links(document, draft.get("internal_links", {}))

        document.save(output_path)
        return str(output_path)

    @staticmethod
    def write_html_artifact(draft: dict, file_stem: str) -> str:
        artifacts_dir = ArtifactWriter._ensure_artifacts_dir()
        output_path = artifacts_dir / f"{slugify(file_stem)}.html"

        metadata = draft.get("metadata", {}) or {}
        sections = draft.get("sections", []) or []
        tables = draft.get("tables", []) or []
        faqs = draft.get("faqs", []) or []
        internal_links = draft.get("internal_links", {}) or {}

        page_title = ArtifactWriter._html_escape(metadata.get("title") or metadata.get("h1") or "Draft")
        h1 = ArtifactWriter._html_escape(metadata.get("h1"))
        meta_description = ArtifactWriter._html_escape(metadata.get("meta_description"))
        intro_snippet = ArtifactWriter._html_escape(metadata.get("intro_snippet"))

        html_parts: list[str] = [
            "<!DOCTYPE html>",
            '<html lang="en">',
            "<head>",
            '  <meta charset="utf-8" />',
            '  <meta name="viewport" content="width=device-width, initial-scale=1" />',
            f"  <title>{page_title}</title>",
            f'  <meta name="description" content="{meta_description}" />',
            "  <style>",
            "    body { font-family: Arial, sans-serif; color: #111827; margin: 40px; line-height: 1.6; }",
            "    h1, h2, h3 { color: #111827; }",
            "    .meta-block, .section-block, .table-block, .faq-block, .links-block { margin-bottom: 28px; }",
            "    .muted { color: #6b7280; }",
            "    table { border-collapse: collapse; width: 100%; margin-top: 12px; }",
            "    th, td { border: 1px solid #d1d5db; padding: 10px; text-align: left; vertical-align: top; }",
            "    th { background: #f3f4f6; }",
            "    ul { margin-top: 8px; }",
            "    li { margin-bottom: 6px; }",
            "    .review-banner { background: #fff3cd; border: 2px solid #c03a00; color: #c03a00; "
            "padding: 14px 18px; border-radius: 6px; font-weight: bold; margin-bottom: 24px; }",
            "  </style>",
            "</head>",
            "<body>",
        ]

        # H1: Inject JSON-LD schema markup (<script type="application/ld+json">) before </head>.
        schema_blocks = SchemaMarkupGenerator.generate_all(draft)
        schema_script_tags = SchemaMarkupGenerator.to_script_tags(schema_blocks)
        if schema_script_tags:
            try:
                head_close_idx = html_parts.index("</head>")
                for tag in reversed(schema_script_tags):
                    html_parts.insert(head_close_idx, tag)
            except ValueError:
                # Fallback: append before <body> if </head> not found
                html_parts.extend(schema_script_tags)

        # E3: Inject review banner at the top when draft failed hard validation.
        if draft.get("needs_review"):
            approval = draft.get("quality_report", {}).get("approval_status", "unknown")
            html_parts.append(
                f'  <div class="review-banner">⚠️ REVIEW REQUIRED — Validation failed '
                f"(approval_status={html.escape(str(approval))}). "
                "Do not publish without editorial review.</div>"
            )

        html_parts.extend([
            f"  <h1>{h1}</h1>",
            '  <div class="meta-block">',
            f"    <p>{intro_snippet}</p>",
            "    <h2>SEO Metadata</h2>",
            f"    <p><strong>Title:</strong> {page_title}</p>",
            f"    <p><strong>Meta Description:</strong> {meta_description}</p>",
            "  </div>",
        ])

        if sections:
            html_parts.append('  <div class="section-block">')
            html_parts.append("    <h2>Editorial Sections</h2>")
            for section in sections:
                title = ArtifactWriter._html_escape(section.get("title"))
                body = str(section.get("body") or "").strip()
                key_points: list[str] = [
                    str(point).strip()
                    for point in (section.get("key_points") or [])
                    if point and str(point).strip()
                ]
                html_parts.append(f"    <h3>{title}</h3>")
                if body:
                    for paragraph in body.split("\n"):
                        paragraph = paragraph.strip()
                        if paragraph:
                            html_parts.append(f"    <p>{html.escape(paragraph)}</p>")
                else:
                    html_parts.append("    <p class='muted'>No grounded narrative available for this section.</p>")
                if key_points:
                    html_parts.append("    <ul>")
                    for point in key_points:
                        html_parts.append(f"      <li>{html.escape(point)}</li>")
                    html_parts.append("    </ul>")
            html_parts.append("  </div>")

        if tables:
            html_parts.append('  <div class="table-block">')
            html_parts.append("    <h2>Key Data Tables</h2>")
            for table_block in tables:
                title = ArtifactWriter._html_escape(table_block.get("title"))
                summary = str(table_block.get("summary") or "").strip()
                columns = table_block.get("columns", []) or []
                rows = table_block.get("rows", []) or []

                html_parts.append(f"    <h3>{title}</h3>")
                if summary:
                    html_parts.append(f"    <p>{html.escape(summary)}</p>")

                if not columns:
                    html_parts.append("    <p class='muted'>No structured columns available.</p>")
                    continue

                html_parts.append("    <table>")
                html_parts.append("      <thead><tr>")
                for column_name in columns:
                    html_parts.append(f"        <th>{html.escape(ArtifactWriter._column_label(column_name))}</th>")
                html_parts.append("      </tr></thead>")
                html_parts.append("      <tbody>")

                if rows:
                    for row in rows:
                        html_parts.append("        <tr>")
                        for column_name in columns:
                            cell_text = OutputFormatter.format_cell(column_name, row.get(column_name))
                            html_parts.append(f"          <td>{html.escape(str(cell_text))}</td>")
                        html_parts.append("        </tr>")
                else:
                    html_parts.append("        <tr>")
                    for _ in columns:
                        html_parts.append("          <td>—</td>")
                    html_parts.append("        </tr>")

                html_parts.append("      </tbody>")
                html_parts.append("    </table>")
            html_parts.append("  </div>")

        if faqs:
            html_parts.append('  <div class="faq-block">')
            html_parts.append("    <h2>Frequently Asked Questions</h2>")
            for faq in faqs:
                question = ArtifactWriter._html_escape(faq.get("question"))
                answer = str(faq.get("answer") or "").strip()
                html_parts.append(f"    <h3>{question}</h3>")
                if answer:
                    for paragraph in answer.split("\n"):
                        paragraph = paragraph.strip()
                        if paragraph:
                            html_parts.append(f"    <p>{html.escape(paragraph)}</p>")
                else:
                    html_parts.append("    <p class='muted'>No grounded answer available.</p>")
            html_parts.append("  </div>")

        if internal_links:
            title_map = {
                "sale_unit_type_links": "Unit Type Links",
                "sale_property_type_links": "Property Type Links",
                "sale_quick_links": "Quick Links",
                "nearby_locality_links": "Nearby Locality Links",
                "top_project_links": "Top Project Links",
                "featured_project_links": "Featured Project Links",
            }
            html_parts.append('  <div class="links-block">')
            html_parts.append("    <h2>Internal Links</h2>")

            for key, group_title in title_map.items():
                group_links = ArtifactWriter._normalize_link_items(internal_links.get(key, []))
                if not group_links:
                    continue

                html_parts.append(f"    <h3>{html.escape(group_title)}</h3>")
                html_parts.append("    <ul>")
                for item in group_links:
                    label = ArtifactWriter._safe_text(
                        item.get("label")
                        or item.get("unitType")
                        or item.get("propertyType")
                        or "Link"
                    )
                    url = ArtifactWriter._safe_text(item.get("url"))
                    if url != "—":
                        html_parts.append(
                            f'      <li><a href="{html.escape(url)}" target="_blank" rel="noreferrer">{html.escape(label)}</a></li>'
                        )
                    else:
                        html_parts.append(f"      <li>{html.escape(label)}</li>")
                html_parts.append("    </ul>")
            html_parts.append("  </div>")

        html_parts.append("</body>")
        html_parts.append("</html>")

        with output_path.open("w", encoding="utf-8") as file:
            file.write("\n".join(html_parts))

        return str(output_path)

    @staticmethod
    def write_blueprint(blueprint: dict) -> str:
        entity = blueprint["entity"]["entity_name"]
        page_type = blueprint["page_type"]
        file_stem = f"{entity}-{page_type}-blueprint"
        return ArtifactWriter.write_json_artifact(blueprint, file_stem)

    @staticmethod
    def write_keyword_intelligence(keyword_intelligence: dict) -> str:
        entity = keyword_intelligence["entity"]["entity_name"]
        page_type = keyword_intelligence["page_type"]
        file_stem = f"{entity}-{page_type}-keyword-intelligence"
        return ArtifactWriter.write_json_artifact(keyword_intelligence, file_stem)

    @staticmethod
    def write_content_plan(content_plan: dict) -> str:
        entity = content_plan["entity"]["entity_name"]
        page_type = content_plan["page_type"]
        file_stem = f"{entity}-{page_type}-content-plan"
        return ArtifactWriter.write_json_artifact(content_plan, file_stem)

    @staticmethod
    def write_draft_bundle(
        draft: dict,
        export_formats: list[str] | None = None,
    ) -> dict[str, str]:
        ArtifactWriter._guard_review_block(draft)

        requested_formats = export_formats or list(settings.draft_default_export_formats)
        requested_formats = [str(item).strip().lower() for item in requested_formats]
        requested_formats = list(dict.fromkeys(requested_formats))

        entity = draft["entity"]["entity_name"]
        page_type = draft["page_type"]
        file_stem = f"{entity}-{page_type}-draft"

        artifact_paths: dict[str, str] = {}

        if "json" in requested_formats:
            artifact_paths["json_path"] = ArtifactWriter.write_json_artifact(draft, file_stem)

        if "markdown" in requested_formats:
            artifact_paths["markdown_path"] = ArtifactWriter.write_markdown_artifact(
                draft["markdown_draft"],
                file_stem,
            )

        if "docx" in requested_formats:
            artifact_paths["docx_path"] = ArtifactWriter.write_docx_artifact(draft, file_stem)

        if "html" in requested_formats:
            artifact_paths["html_path"] = ArtifactWriter.write_html_artifact(draft, file_stem)

        return artifact_paths